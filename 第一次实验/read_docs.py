from docx import Document

# 读取实验要求文档
try:
    doc1 = Document('17782065361055.doc')
    print("=" * 80)
    print("实验要求文档 (17782065361055.doc):")
    print("=" * 80)
    for i, para in enumerate(doc1.paragraphs):
        if para.text.strip():
            print(f"段落 {i+1}: {para.text}")
    print("\n")
except Exception as e:
    print(f"读取 17782065361055.doc 失败：{e}")

# 读取实验报告文档
try:
    doc2 = Document('17782065431497.docx')
    print("=" * 80)
    print("实验报告文档 (17782065431497.docx):")
    print("=" * 80)
    for i, para in enumerate(doc2.paragraphs):
        if para.text.strip():
            print(f"段落 {i+1}: {para.text}")
except Exception as e:
    print(f"读取 17782065431497.docx 失败：{e}")
